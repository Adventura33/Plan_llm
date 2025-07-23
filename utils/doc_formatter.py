# utils/doc_formatter.py

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn # Необходим, если придется использовать низкоуровневые настройки границ
from docx.oxml import OxmlElement # Необходим, если придется использовать низкоуровневые настройки границ


# Функция для отладки: читаем сохраненный документ и выводим содержимое таблицы
def debug_read_docx_table(file_path):
    try:
        doc = Document(file_path)
        print(f"\n--- Отладка содержимого сохраненного DOCX: {file_path} ---")
        tables = doc.tables
        if not tables:
            print("В документе нет таблиц.")
            return

        table = tables[0]
        print(f"Найдена таблица с {len(table.rows)} строками и {len(table.columns)} колонками.")

        header_cells = table.rows[0].cells
        srok_col_idx = -1
        ispolnitel_col_idx = -1
        for i, cell in enumerate(header_cells):
            if cell.text.strip() == 'Срок':
                srok_col_idx = i
            if cell.text.strip() == 'Исполнитель':
                ispolnitel_col_idx = i
            
        if srok_col_idx == -1:
            print("Колонка 'Срок' не найдена в заголовках таблицы.")
        if ispolnitel_col_idx == -1:
            print("Колонка 'Исполнитель' не найдена в заголовках таблицы.")

        print(f"Содержимое колонки 'Срок' (первые 5 строк):")
        for i, row in enumerate(table.rows):
            if i == 0: # Пропускаем заголовок
                continue
            if i > 5: # Выводим только первые 5 строк для краткости
                break
            try:
                cell_text_srok = row.cells[srok_col_idx].text.strip()
                cell_text_ispolnitel = row.cells[ispolnitel_col_idx].text.strip()
                print(f"  Строка {i} - Срок: '{cell_text_srok}', Исполнитель: '{cell_text_ispolnitel}'")
            except IndexError:
                print(f"  Строка {i}: (ошибка доступа к ячейке)")
        print("--- Конец отладки содержимого DOCX ---")
    except Exception as e:
        print(f"Ошибка при чтении сохраненного DOCX для отладки: {e}")


def create_investigation_plan_doc(plan_data, output_path):
    """
    Создает документ Word (.docx) со сгенерированным планом расследования,
    ориентируясь на формат нового образца.
    """
    document = Document()

    # Установка полей страницы (примерные значения для более плотного текста)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Доступная ширина страницы (Letter, 8.5 дюймов - 0.7*2 поля = 7.1 дюймов)
    available_width = Inches(7.1)

    # ОБРАЗЕЦ
    p = document.add_paragraph()
    p.add_run('ОБРАЗЕЦ').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # УТВЕРЖДАЮ
    p = document.add_paragraph()
    p.add_run('«УТВЕРЖДАЮ»').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = document.add_paragraph()
    p.add_run('Заместитель руководитель ДЭР').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Пустые строки для подписи
    document.add_paragraph('____________________________').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('____________________________').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph('«_______» ____________________ 2023 г.').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Заголовок ПЛАН
    document.add_paragraph() # Пустая строка
    p = document.add_paragraph()
    p.add_run('ПЛАН').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Детализация плана
    plan_title_info = plan_data.get("plan_title_info", {})
    plan_text = document.add_paragraph(
        f"параллельного финансового расследования уголовного дела №{plan_title_info.get('номер_дела', '_______')} "
        f"по факту {plan_title_info.get('факт', '___________________________')} "
        f"по ст. {plan_title_info.get('статья_ук_рк', '_______ УК РК.')}."
    )
    plan_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"{plan_title_info.get('год', '_______')} г.").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_paragraph() # Пустая строка перед таблицей

    # --- ИЗМЕНЕНИЕ ЗДЕСЬ: ДОБАВЛЕНИЕ СТИЛЯ ТАБЛИЦЫ ---
    # 'Table Grid' - это встроенный стиль, который обычно имеет видимые границы.
    table = document.add_table(rows=1, cols=4, style='Table Grid') 
    table.autofit = False
    table.allow_autofit = False # Добавлено для жесткого контроля ширины

    # Настройка ширины колонок (остается как есть)
    col_width_np = Inches(0.4)   # № п.п. (очень узкая)
    col_width_srok = Inches(1.0) # Срок (визуально узкая)
    col_width_ispolnitel = Inches(2.3) # Исполнитель (достаточно широкая для должности и ФИО)
    
    # Оставшаяся ширина для "Следственные действия"
    col_width_deystviya = available_width - col_width_np - col_width_srok - col_width_ispolnitel

    table.columns[0].width = col_width_np
    table.columns[1].width = col_width_deystviya # Следственные действия
    table.columns[2].width = col_width_ispolnitel
    table.columns[3].width = col_width_srok

    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п.п.'
    hdr_cells[1].text = 'Следственные действия'
    hdr_cells[2].text = 'Исполнитель'
    hdr_cells[3].text = 'Срок'

    # Применение жирного шрифта и выравнивания к заголовкам
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Заполнение таблицы данными
    for action in plan_data.get("actions", []):
        row_cells = table.add_row().cells
        row_cells[0].text = str(action.get("номер", ""))
        row_cells[1].text = action.get("действие", "")
        row_cells[2].text = action.get("исполнитель", "")
        row_cells[3].text = action.get("срок", "")
        
        for i, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:
                # Номер по центру, остальные по левому краю
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                # Выравнивание по верхнему краю ячейки
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Добавление заключительных параграфов
    document.add_paragraph()
    document.add_paragraph(
        "Следственное действие должно раскрывать цель его проведения, участников и сроки исполнения, при этом, "
        "следственные мероприятия необходимо группировать по направлениям исследований."
    )
    document.add_paragraph(
        "При этом, каждое планируемое следственное действие и мероприятие должны быть отражены раздельно."
    )
    document.add_paragraph()
    document.add_paragraph(
        "Данные следственные действия не являются исчерпывающими, по мере необходимости провести и другие "
        "следственно-оперативные мероприятия с составлением дополнительного плана."
    )
    
    # Подписи в конце
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run('Следователь').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = document.add_paragraph('____________________________').alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run('«Согласовано»').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = document.add_paragraph()
    p.add_run('Руководитель СУ').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph('____________________________').alignment = WD_ALIGN_PARAGRAPH.LEFT


    document.save(output_path)
    print(f"Документ плана расследования сохранен по пути: {output_path}")
    
    # --- НОВОЕ: Отладочный вызов чтения DOCX ---
    debug_read_docx_table(output_path)
    # --- КОНЕЦ НОВОГО ---
